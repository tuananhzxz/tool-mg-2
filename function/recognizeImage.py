import os
import pytesseract
from PIL import Image
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import threading
import re


class OptimizedMangaOCR:
    def __init__(self, root):
        self.root = root
        self.root.title("Manga OCR - Tối ưu cho ảnh lớn")
        self.root.geometry("700x500")

        self.selected_files = []
        self.progress_var = tk.DoubleVar()
        self.status_var = tk.StringVar()
        self.status_var.set("Sẵn sàng")

        self.create_widgets()

    def create_widgets(self):
        # Frame chính
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Nút chọn ảnh
        file_frame = ttk.Frame(main_frame)
        file_frame.pack(fill=tk.X, pady=5)

        ttk.Button(file_frame, text="Chọn ảnh truyện tranh", command=self.select_files).pack(side=tk.LEFT, padx=5)
        ttk.Button(file_frame, text="Xóa tất cả", command=self.clear_files).pack(side=tk.LEFT, padx=5)

        # Danh sách file
        list_frame = ttk.LabelFrame(main_frame, text="Danh sách ảnh")
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.file_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set)
        self.file_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.config(command=self.file_listbox.yview)

        # Tùy chọn
        options_frame = ttk.LabelFrame(main_frame, text="Tùy chọn OCR")
        options_frame.pack(fill=tk.X, pady=5)

        # Ngôn ngữ
        ttk.Label(options_frame, text="Ngôn ngữ:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.lang_var = tk.StringVar(value="eng")
        lang_combo = ttk.Combobox(options_frame, textvariable=self.lang_var, values=["eng", "vie", "eng+vie"])
        lang_combo.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)

        # Threshold cho phân đoạn
        ttk.Label(options_frame, text="Khoảng cách giữa các đoạn (pixel):").grid(row=1, column=0, padx=5, pady=5,
                                                                                 sticky=tk.W)
        self.gap_threshold_var = tk.IntVar(value=20)
        gap_entry = ttk.Entry(options_frame, textvariable=self.gap_threshold_var, width=5)
        gap_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)

        # Tùy chọn tiền xử lý
        self.preprocess_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Tiền xử lý ảnh (tăng độ chính xác, chậm hơn)",
                        variable=self.preprocess_var).grid(row=2, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W)

        # Tùy chọn chia nhỏ ảnh
        self.split_image_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="Chia nhỏ ảnh lớn để xử lý (tăng tốc với ảnh lớn)",
                        variable=self.split_image_var).grid(row=3, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W)

        # Chiều cao mỗi phần khi chia nhỏ
        ttk.Label(options_frame, text="Chiều cao mỗi phần khi chia (pixel):").grid(row=4, column=0, padx=5, pady=5,
                                                                                   sticky=tk.W)
        self.chunk_height_var = tk.IntVar(value=1000)
        chunk_entry = ttk.Entry(options_frame, textvariable=self.chunk_height_var, width=5)
        chunk_entry.grid(row=4, column=1, padx=5, pady=5, sticky=tk.W)

        # Nút xử lý
        process_frame = ttk.Frame(main_frame)
        process_frame.pack(fill=tk.X, pady=10)

        ttk.Button(process_frame, text="Xử lý OCR và lưu vào Word",
                   command=self.start_processing).pack(side=tk.LEFT, padx=5)

        # Thanh tiến trình
        ttk.Label(main_frame, text="Tiến trình:").pack(anchor=tk.W, padx=5)
        self.progress_bar = ttk.Progressbar(main_frame, orient=tk.HORIZONTAL,
                                            length=100, mode='determinate', variable=self.progress_var)
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)

        # Thanh trạng thái
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(fill=tk.X, side=tk.BOTTOM, pady=5)

    def select_files(self):
        filetypes = [("Image files", "*.jpg *.jpeg *.png *.webp")]
        files = filedialog.askopenfilenames(title="Chọn ảnh truyện tranh", filetypes=filetypes)

        if files:
            self.selected_files.extend(files)
            self.update_file_list()

    def update_file_list(self):
        self.file_listbox.delete(0, tk.END)
        for i, file in enumerate(self.selected_files, 1):
            self.file_listbox.insert(tk.END, f"{i}. {os.path.basename(file)}")

        self.status_var.set(f"Đã chọn {len(self.selected_files)} ảnh")

    def clear_files(self):
        self.selected_files = []
        self.file_listbox.delete(0, tk.END)
        self.status_var.set("Danh sách ảnh đã được xóa")

    def preprocess_image(self, image):
        """Tiền xử lý ảnh nâng cao để cải thiện độ chính xác OCR"""
        # Chuyển sang grayscale
        if image.mode != 'L':
            image = image.convert('L')

        # Loại bỏ nhiễu với Gaussian blur nhẹ
        from PIL import ImageFilter
        image = image.filter(ImageFilter.GaussianBlur(radius=0.5))

        # Tăng cường tương phản
        from PIL import ImageEnhance
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.5)

        # Nhị phân hóa ảnh (thresholding) - tách rõ chữ và nền
        threshold = 180  # Điều chỉnh ngưỡng này theo nhu cầu
        from PIL import ImageOps
        image = ImageOps.invert(image)  # Đảo ngược để chữ là màu đen trên nền trắng
        image = image.point(lambda x: 255 if x > threshold else 0)

        return image

    def split_image_to_chunks(self, image):
        """Chia ảnh lớn thành các phần nhỏ hơn để xử lý"""
        width, height = image.size
        chunk_height = self.chunk_height_var.get()

        # Tính số lượng phần cần chia
        num_chunks = max(1, height // chunk_height + (1 if height % chunk_height > 0 else 0))

        chunks = []
        for i in range(num_chunks):
            top = i * chunk_height
            bottom = min((i + 1) * chunk_height, height)

            chunk = image.crop((0, top, width, bottom))
            chunks.append(chunk)

        return chunks

    def extract_text_from_image(self, image_path):
        try:
            # Đọc ảnh
            image = Image.open(image_path)

            # Tiền xử lý
            if self.preprocess_var.get():
                image = self.preprocess_image(image)

            # Cấu hình nâng cao cho Tesseract
            custom_config = r'--oem 1 --psm 6 -c preserve_interword_spaces=1'

            # Thêm điều kiện phát hiện dọc nếu là manga tiếng Nhật
            if self.lang_var.get() == "jpn":
                custom_config += ' -c textord_orientation=1'

            # Xử lý theo từng phần hoặc toàn bộ
            if self.split_image_var.get() and image.height > self.chunk_height_var.get():
                chunks = self.split_image_to_chunks(image)
                all_text = ""

                for chunk in chunks:
                    chunk_text = pytesseract.image_to_string(
                        chunk,
                        lang=self.lang_var.get(),
                        config=custom_config
                    )
                    all_text += chunk_text + "\n"
            else:
                all_text = pytesseract.image_to_string(
                    image,
                    lang=self.lang_var.get(),
                    config=custom_config
                )

            # Phân tách hội thoại
            paragraphs = self.separate_dialogues(all_text)

            return paragraphs
        except Exception as e:
            return [f"Lỗi khi xử lý ảnh: {str(e)}"]

    def separate_dialogues(self, text):
        """Phân tách văn bản thành các đoạn hội thoại với xử lý nâng cao"""
        # Loại bỏ ký tự rác thường gặp trong OCR
        text = re.sub(r'[^\w\s\.,!?;:\'"-]', '', text)

        # Loại bỏ dòng trống liên tiếp
        text = re.sub(r'\n\s*\n', '\n\n', text)

        # Gom nhóm các dòng gần nhau (khoảng cách nhỏ)
        lines = text.split('\n')
        gap_threshold = self.gap_threshold_var.get()

        grouped_lines = []
        current_group = []

        for line in lines:
            line = line.strip()
            if not line:  # Dòng trống
                if current_group:
                    grouped_lines.append(' '.join(current_group))
                    current_group = []
            else:
                current_group.append(line)

        # Thêm nhóm cuối cùng nếu có
        if current_group:
            grouped_lines.append(' '.join(current_group))

        # Loại bỏ nhóm quá ngắn hoặc không có ý nghĩa
        final_paragraphs = []
        for para in grouped_lines:
            # Loại bỏ khoảng trắng thừa
            cleaned = re.sub(r'\s+', ' ', para).strip()

            # Kiểm tra độ dài tối thiểu và tỷ lệ ký tự hợp lệ
            if len(cleaned) >= 3:
                # Tính tỷ lệ ký tự hợp lệ
                valid_chars = sum(1 for c in cleaned if c.isalnum() or c in '.,!?;: ')
                valid_ratio = valid_chars / len(cleaned) if len(cleaned) > 0 else 0

                if valid_ratio > 0.7:  # Ngưỡng có thể điều chỉnh
                    final_paragraphs.append(cleaned)

        return final_paragraphs

    def start_processing(self):
        """Bắt đầu quá trình xử lý trong một thread riêng biệt"""
        if not self.selected_files:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất một ảnh để xử lý")
            return

        # Vô hiệu hóa các nút trong khi xử lý
        for widget in self.root.winfo_children():
            if isinstance(widget, ttk.Button):
                widget.configure(state=tk.DISABLED)

        # Đặt lại thanh tiến trình
        self.progress_var.set(0)

        # Bắt đầu thread xử lý
        thread = threading.Thread(target=self.process_images)
        thread.daemon = True
        thread.start()

    def edit_results(self, all_results, doc):
        """Cho phép người dùng hiệu chỉnh kết quả OCR trước khi lưu"""
        edit_window = tk.Toplevel(self.root)
        edit_window.title("Hiệu chỉnh kết quả OCR")
        edit_window.geometry("800x600")

        # Tạo notebook để hiển thị từng trang riêng biệt
        notebook = ttk.Notebook(edit_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_widgets = []

        # Tạo tab cho mỗi trang ảnh
        for result in all_results:
            frame = ttk.Frame(notebook)
            notebook.add(frame, text=result["filename"])

            # Text widget để chỉnh sửa
            text_widget = tk.Text(frame, wrap=tk.WORD, width=80, height=20)
            text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

            # Hiển thị kết quả OCR
            for i, paragraph in enumerate(result['paragraphs'], 1):
                text_widget.insert(tk.END, f"{i}. {paragraph}\n\n")

            text_widgets.append(text_widget)

        # Nút lưu
        def save_edited():
            # Cập nhật kết quả từ text widgets
            for i, text_widget in enumerate(text_widgets):
                content = text_widget.get("1.0", tk.END)

                # Phân tách lại thành từng đoạn
                paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

                # Loại bỏ số thứ tự nếu có
                cleaned_paragraphs = []
                for p in paragraphs:
                    if p and p[0].isdigit() and ". " in p:
                        p = p.split(". ", 1)[1]
                    cleaned_paragraphs.append(p)

                all_results[i]['paragraphs'] = cleaned_paragraphs

            # Thêm vào document
            for result in all_results:
                doc.add_heading(f'Trang: {result["filename"]}', level=1)
                for j, paragraph in enumerate(result['paragraphs'], 1):
                    doc.add_paragraph(f"{j}. {paragraph}")
                doc.add_paragraph('-----------------------------------')

            # Lưu document
            self.save_document(doc)
            edit_window.destroy()

        save_button = ttk.Button(edit_window, text="Lưu thay đổi", command=save_edited)
        save_button.pack(pady=10)

    def save_document(self, doc):
        """Lưu document Word với kết quả OCR"""
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Lưu kết quả OCR"
        )

        if save_path:
            try:
                doc.save(save_path)
                self.status_var.set(f"Đã lưu kết quả OCR vào {save_path}")
                messagebox.showinfo("Thành công", f"Đã lưu kết quả OCR vào:\n{save_path}")
            except Exception as e:
                self.status_var.set(f"Lỗi khi lưu: {str(e)}")
                messagebox.showerror("Lỗi", f"Không thể lưu tài liệu: {str(e)}")
        else:
            self.status_var.set("Đã hủy lưu tài liệu")

    def process_images(self):
        """Xử lý OCR cho tất cả ảnh với khả năng hiệu chỉnh"""
        try:
            doc = Document()
            doc.add_heading('Kết quả OCR từ truyện tranh', 0)

            # Lưu tạm kết quả để hiệu chỉnh
            all_results = []

            total_files = len(self.selected_files)
            for i, file_path in enumerate(self.selected_files):
                filename = os.path.basename(file_path)
                self.status_var.set(f"Đang xử lý ảnh {i + 1}/{total_files}: {filename}")

                # Trích xuất văn bản
                paragraphs = self.extract_text_from_image(file_path)
                all_results.append({
                    'filename': filename,
                    'paragraphs': paragraphs
                })

                # Cập nhật tiến trình
                progress = (i + 1) / total_files * 100
                self.progress_var.set(progress)
                self.root.update_idletasks()

            # Tạo cửa sổ hiệu chỉnh nếu cần
            if messagebox.askyesno("Hiệu chỉnh", "Bạn có muốn hiệu chỉnh kết quả OCR trước khi lưu không?"):
                self.edit_results(all_results, doc)
            else:
                # Thêm trực tiếp vào document
                for result in all_results:
                    doc.add_heading(f'Trang: {result["filename"]}', level=1)
                    for j, paragraph in enumerate(result['paragraphs'], 1):
                        doc.add_paragraph(f"{j}. {paragraph}")
                    doc.add_paragraph('-----------------------------------')

                # Lưu document
                self.save_document(doc)

        except Exception as e:
            self.status_var.set(f"Lỗi: {str(e)}")
            messagebox.showerror("Lỗi", str(e))
        finally:
            # Kích hoạt lại các nút
            for widget in self.root.winfo_children():
                if isinstance(widget, ttk.Button):
                    widget.configure(state=tk.NORMAL)


if __name__ == "__main__":
    try:
        import pytesseract
        from PIL import Image
        from docx import Document
    except ImportError as e:
        print(f"Thiếu thư viện: {e}")
        print("Vui lòng cài đặt các thư viện cần thiết bằng lệnh:")
        print("pip install pytesseract pillow python-docx")
    else:
        root = tk.Tk()
        app = OptimizedMangaOCR(root)
        root.mainloop()