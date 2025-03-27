import time
from flask import Flask, render_template, request, jsonify, send_file
import os
from werkzeug.utils import secure_filename
from PIL import Image, ImageEnhance
import shutil
from function.changeImage import convert_images
from function.cutmergeimage import ImageProcessor
from function.mergeWord import merge_word_documents
from function.renameImage import rename_files
import requests
from bs4 import BeautifulSoup
import io
import zipfile
import re
from flask_cors import CORS
from function.ocr_processor import process_folder
from function.speechbubble import SpeechBubbleProcessor
from docx import Document
from google import genai
from google.genai import types
from dotenv import load_dotenv
from service.dowloadImg import (
    download_selected_images,
)

# Load environment variables
load_dotenv()

app = Flask(__name__, static_folder='public')
CORS(app)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max-limit

# Đảm bảo thư mục uploads tồn tại
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


# Hàm xử lý lỗi chung
def handle_error(e):
    error_msg = str(e)
    if isinstance(e, requests.exceptions.RequestException):
        error_msg = "Lỗi kết nối: Không thể tải dữ liệu từ URL"
    elif isinstance(e, IOError):
        error_msg = "Lỗi đọc/ghi file"
    elif isinstance(e, Image.UnidentifiedImageError):
        error_msg = "File ảnh không hợp lệ hoặc bị hỏng"
    return jsonify({'error': error_msg})

@app.route('/')
def index():
    functions = {
        'changeImage': 'Chuyển đổi định dạng ảnh',
        'cutmergeimage': 'Cắt và ghép ảnh',
        'dowloaf': 'Tải ảnh từ web',
        'mergeWord': 'Gộp file Word',
        'renameImage': 'Đổi tên ảnh',
        'ocr': 'Nhận dạng chữ trong ảnh (OCR)'
    }
    return render_template('index.html', functions=functions)

@app.route('/function/<function_name>')
def function_page(function_name):
    templates = {
        'changeImage': 'functions/changeimage.html',
        'cutmergeimage': 'functions/cutmergeimage.html',
        'dowloaf': 'functions/dowloaf.html',
        'mergeWord': 'functions/mergeword.html',
        'renameImage': 'functions/renameimage.html',
        'ocr': 'functions/ocr.html',
        'translate': 'functions/translate.html',
        'remove_speech_bubbles': 'functions/speechbubble.html'
    }
    
    if function_name not in templates:
        return jsonify({'error': 'Chức năng không tồn tại'}), 404
        
    return render_template(templates[function_name])

@app.route('/execute/changeImage', methods=['POST'])
def execute_change_image():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'Không có file được tải lên'})
        
        files = request.files.getlist('files')
        target_format = request.form.get('target_format', 'webp')
        
        if not files or files[0].filename == '':
            return jsonify({'error': 'Không có file được chọn'})
            
        # Tạo thư mục tạm thời
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_convert')
        output_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'output_convert')
        
        for dir_path in [temp_dir, output_dir]:
            if os.path.exists(dir_path):
                shutil.rmtree(dir_path)
            os.makedirs(dir_path)
            
        # Lưu và xử lý từng file
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
                
                # Xác định định dạng nguồn và chuyển đổi
                with Image.open(file_path) as img:
                    source_format = img.format
                    # Chuyển đổi và lưu với định dạng mới
                    output_path = os.path.join(output_dir, os.path.splitext(filename)[0] + '.' + target_format.lower())
                    if target_format.upper() == 'WEBP':
                        img.save(output_path, 'WEBP', quality=90)
                    elif target_format.upper() == 'JPEG':
                        if img.mode in ('RGBA', 'LA'):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            background.paste(img, mask=img.split()[-1])
                            background.save(output_path, 'JPEG', quality=95)
                        else:
                            img.save(output_path, 'JPEG', quality=95)
                    else:  # PNG
                        img.save(output_path, 'PNG')
        
        # Tạo file zip
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'converted_images.zip')
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(output_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, output_dir)
                    zipf.write(file_path, arcname)
        
        # Xóa thư mục tạm
        shutil.rmtree(temp_dir)
        shutil.rmtree(output_dir)
        
        return jsonify({
            'message': 'Chuyển đổi thành công!',
            'output_files': ['converted_images.zip']
        })
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/cutmergeimage', methods=['POST'])
def execute_cut_merge_image():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    action = request.form.get('action')
    min_height = int(request.form.get('min_height', 2500))  # Chiều cao tối thiểu mặc định 2500px
    images_per_group = int(request.form.get('parts', 2))  # Số ảnh mỗi nhóm khi ghép
    width = request.form.get('width')
    height = request.form.get('height')
    
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
        
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_process')
    output_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'output_process')
    
    for dir_path in [temp_dir, output_dir]:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
        os.makedirs(dir_path)
        
    try:
        # Lưu các file
        saved_files = []
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
                saved_files.append(file_path)
                
                # Resize ảnh nếu có yêu cầu
                if width or height:
                    with Image.open(file_path) as img:
                        new_width = int(width) if width else img.width
                        new_height = int(height) if height else img.height
                        resized = img.resize((new_width, new_height))
                        resized.save(file_path)
        
        processor = ImageProcessor(temp_dir)
        
        try:
            if action == 'split':
                result_files = processor.split_images(min_height)
            else:  # merge
                if len(saved_files) < images_per_group:
                    return jsonify({'error': f'Số lượng ảnh ({len(saved_files)}) phải lớn hơn hoặc bằng số ảnh mỗi nhóm ({images_per_group})'})
                result_files = processor.combine_images(images_per_group)
            
            if not result_files:
                return jsonify({'error': 'Không thể xử lý ảnh. Vui lòng kiểm tra lại các tham số.'})
            
            # Copy các file kết quả vào thư mục output với tên mới
            for i, file_path in enumerate(result_files, 1):
                file_ext = os.path.splitext(file_path)[1]
                new_name = f"{i}{file_ext}"
                new_path = os.path.join(output_dir, new_name)
                shutil.copy2(file_path, new_path)
            
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'processed_images.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, output_dir)
                        zipf.write(file_path, arcname)
            
            # Xóa thư mục tạm
            shutil.rmtree(temp_dir)
            shutil.rmtree(output_dir)
            
            return jsonify({
                'message': 'Xử lý thành công!',
                'output_files': ['processed_images.zip']
            })
            
        except Exception as e:
            return jsonify({'error': str(e)})
            
    except Exception as e:
        return handle_error(e)

@app.route('/execute/mergeWord', methods=['POST'])
def execute_merge_word():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_word')
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    try:
        # Lưu các file
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
        
        # Gộp file
        output_path = os.path.join(temp_dir, 'merged_document.docx')
        success = merge_word_documents(temp_dir, output_path)
        
        if success:
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'merged_documents.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                zipf.write(output_path, 'merged_document.docx')
            
            # Xóa thư mục tạm
            shutil.rmtree(temp_dir)
            
            return jsonify({
                'message': 'Gộp file thành công!',
                'output_files': ['merged_documents.zip']
            })
        else:
            return jsonify({'error': 'Không thể gộp các file'})
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/renameImage', methods=['POST'])
def execute_rename_image():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_rename')
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    try:
        # Lưu các file
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
        
        # Đổi tên file
        rename_files(temp_dir)
        
        # Tạo file zip chứa kết quả
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'renamed_images.zip')
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        
        # Xóa thư mục tạm
        shutil.rmtree(temp_dir)
        
        return jsonify({
            'message': 'Đổi tên file thành công!',
            'output_files': ['renamed_images.zip']
        })
        
    except Exception as e:
        return handle_error(e)

@app.route('/execute/ocr', methods=['POST'])
def execute_ocr():
    try:
        files = request.files.getlist('files')
        api_key = request.form.get('api_key')
        mode = request.form.get('mode', 'ocr')
        genres = request.form.getlist('genres[]')
        styles = request.form.getlist('styles[]')
        
        if not files or not api_key:
            return jsonify({'error': 'Vui lòng cung cấp file và API key'})
        
        # Tạo một tài liệu Word mới
        doc = Document()
        doc.add_heading('Cám ơn đã sử dụng tôi iu bạn có thể mua ủng hộ ly cafe để web phát triển hơn...', 0)
        
        # Thêm thông tin về thể loại và phong cách
        if genres or styles:
            doc.add_paragraph('=== Thông tin truyện ===')
            if genres:
                doc.add_paragraph('Thể loại: ' + ', '.join(genres))
            if styles:
                doc.add_paragraph('Phong cách: ' + ', '.join(styles))
            doc.add_paragraph('---')
        
        extracted_texts = []
        processed_files = []
        failed_files = []
        
        # Xử lý từng file một
        for file in files:
            if not file.filename:
                continue
                
            try:
                # Kiểm tra định dạng file
                file_ext = os.path.splitext(file.filename)[1].lower()
                
                if file_ext in ['.docx', '.doc']:
                    # Xử lý file Word
                    success = process_word_file(file, doc, mode, genres, styles, api_key)
                    if success:
                        processed_files.append(file.filename)
                    else:
                        failed_files.append(file.filename)
                        
                elif file_ext in ['.jpg', '.jpeg', '.png', '.webp']:
                    # Xử lý file ảnh
                    success = process_image_file(file, doc, mode, genres, styles, api_key)
                    if success:
                        processed_files.append(file.filename)
                    else:
                        failed_files.append(file.filename)
                else:
                    failed_files.append(file.filename)
                    continue
                    
            except Exception as e:
                print(f"Lỗi khi xử lý file {file.filename}: {str(e)}")
                failed_files.append(file.filename)
                continue
        
        if not processed_files:
            return jsonify({'error': 'Không thể xử lý bất kỳ file nào'})
        
        # Lưu tài liệu
        output_filename = f'VBCĐ_{int(time.time())}.docx'
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        doc.save(output_path)
        
        return jsonify({
            'success': True,
            'word_files': {'all': [output_filename]},
            'processed_files': processed_files,
            'failed_files': failed_files,
            'extracted_texts': extracted_texts
        })
        
    except Exception as e:
        return jsonify({'error': str(e)})

def process_word_file(file, doc, mode, genres, styles, api_key):
    try:
        docx_doc = Document(file)
        text = ''
        for para in docx_doc.paragraphs:
            text += para.text + '\n'
        
        if text.strip():
            if mode == 'translate':
                # Dịch văn bản
                target_langs = request.form.getlist('target_langs[]')
                for lang in target_langs:
                    prompt = f"""
                        Hãy dịch đoạn văn bản sau sang {getLangName(lang)}.
                        Thể loại: {', '.join(genres) if genres else 'Không xác định'}
                        Phong cách: {', '.join(styles) if styles else 'Không xác định'}
                        
                        Văn bản cần dịch:
                        {text}
                    """
                    client = genai.Client(api_key=api_key)
                    
                    response = client.models.generate_content(
                        model="gemini-2.0-flash",
                        contents=prompt,
                        config={"temperature": 0.2}
                    )
                    
                    if response.text:
                        doc.add_paragraph(f"=== {getLangName(lang)} ===")
                        doc.add_paragraph(response.text)
                        doc.add_paragraph('---')
            else:
                doc.add_paragraph(text)
                doc.add_paragraph('---')
        return True
    except Exception as e:
        print(f"Lỗi khi xử lý file Word {file.filename}: {str(e)}")
        return False

def process_image_file(file, doc, mode, genres, styles, api_key):
    try:
        # Đọc và xử lý ảnh
        image_bytes = file.read()
        img = Image.open(io.BytesIO(image_bytes))
        
        # Kiểm tra kích thước ảnh
        max_size = (1920, 1080)
        if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        # Chuyển sang grayscale và tăng độ tương phản
        img = img.convert('L')
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)
        
        # Chuyển ảnh đã xử lý thành bytes
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        # Tạo prompt cho việc nhận dạng văn bản
        prompt = f"""
            NHIỆM VỤ: Nhận dạng và trích xuất văn bản từ ảnh với độ chính xác cao nhất.
            
            YÊU CẦU CHẤT LƯỢNG:
            1. Nhận dạng chính xác 100% nội dung văn bản, kể cả chữ nhỏ
            2. Phân biệt rõ các đoạn văn bản khác nhau, các bóng thoại khác nhau
            3. Giữ nguyên vị trí và thứ tự của các bóng thoại
            4. Không bỏ sót bất kỳ ký tự nào
            
            QUY TẮC XỬ LÝ:
            1. Loại bỏ các yếu tố không phải văn bản 
            2. QUAN TRỌNG: Xử lý mỗi bóng thoại (speech bubble) như MỘT CÂU HOÀN CHỈNH TRÊN MỘT DÒNG DUY NHẤT
            3. Mỗi bóng thoại riêng biệt sẽ được xuất ra thành một dòng văn bản riêng biệt
            4. Giữ nguyên các dấu câu và định dạng đặc biệt
            
            ĐỊNH DẠNG ĐẦU RA:
            - Mỗi bóng thoại trên một dòng riêng
            - Giữ nguyên các dấu câu và định dạng
            - Không thêm bất kỳ chú thích hay giải thích nào
            - Không tách văn bản trong một bóng thoại thành nhiều dòng
        """
        
        client = genai.Client(api_key=api_key)
        
        # Gửi yêu cầu đến Gemini API với cấu hình tối ưu
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=[
                {"text": prompt},
                {"inline_data": {"mime_type": "image/png", "data": img_byte_arr}}
            ],
            config={
                "temperature": 0.1,  # Giảm temperature để tăng độ chính xác
                "max_output_tokens": 2048,  # Tăng max tokens để xử lý văn bản dài
                "top_p": 0.8,
                "top_k": 40,
                "candidate_count": 1
            }
        )
        
        if response.text:
            # Thêm thông tin về file đang xử lý
            doc.add_paragraph(f"=== File: {file.filename} ===")
            
            if mode == 'translate':
                # Dịch văn bản
                target_langs = request.form.getlist('target_langs[]')
                for lang in target_langs:
                    prompt = f"""
                        Hãy dịch đoạn văn bản sau sang {getLangName(lang)}.
                        Thể loại: {', '.join(genres) if genres else 'Không xác định'}
                        Phong cách: {', '.join(styles) if styles else 'Không xác định'}
                        
                        Văn bản cần dịch:
                        {response.text}
                    """
                    
                    response = client.models.generate_content(
                        model="gemini-2.0-flash",
                        contents=prompt,
                        config={
                            "temperature": 0.2,
                            "max_output_tokens": 2048
                        }
                    )
                    
                    if response.text:
                        doc.add_paragraph(f"=== {getLangName(lang)} ===")
                        doc.add_paragraph(response.text)
                        doc.add_paragraph('---')
            else:
                doc.add_paragraph(response.text)
                doc.add_paragraph('---')
                
            # Thêm delay giữa các lần xử lý để tránh rate limit
            time.sleep(1)
            
        return True
    except Exception as e:
        print(f"Lỗi khi xử lý file ảnh {file.filename}: {str(e)}")
        return False

# Hàm chuyển đổi mã ngôn ngữ thành tên
def getLangName(code):
    langs = {
        'vie': 'Tiếng Việt',
        'eng': 'Tiếng Anh',
        'jpn': 'Tiếng Nhật',
        'kor': 'Tiếng Hàn'
    }
    return langs.get(code, code)

@app.route('/execute/merge_ocr', methods=['POST'])
def execute_merge_ocr():
    if 'files' not in request.files:
        return jsonify({'error': 'Không có file được tải lên'})
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({'error': 'Không có file được chọn'})
    
    # Tạo thư mục tạm thời
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_merge_ocr')
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)
    
    try:
        # Lưu các file Word
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_path = os.path.join(temp_dir, filename)
                file.save(file_path)
        
        # Gộp các file Word
        output_path = os.path.join(temp_dir, 'merged_ocr.docx')
        success = merge_word_documents(temp_dir, output_path)
        
        if success:
            # Tạo file zip chứa kết quả
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'merged_ocr.zip')
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                zipf.write(output_path, 'merged_ocr.docx')
            
            # Xóa thư mục tạm
            shutil.rmtree(temp_dir)
            
            return jsonify({
                'message': 'Gộp file thành công!',
                'output_files': ['merged_ocr.zip']
            })
        else:
            return jsonify({'error': 'Không thể gộp các file'})
        
    except Exception as e:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        return handle_error(e)

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            return jsonify({'error': 'File không tồn tại'})
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return handle_error(e)

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File quá lớn. Kích thước tối đa là 16MB'}), 413

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Lỗi máy chủ nội bộ'}), 500

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Không tìm thấy trang'}), 404

@app.route('/remove_speech_bubbles', methods=['GET', 'POST'])
def process_speech_bubbles():
    if request.method == 'GET':
        return render_template('functions/speechbubble.html')
    
    try:
        # Kiểm tra file upload
        if 'files' not in request.files:
            return jsonify({'error': 'Không tìm thấy file nào được tải lên'}), 400
            
        files = request.files.getlist('files')
        if not files or files[0].filename == '':
            return jsonify({'error': 'Không có file nào được chọn'}), 400
            
        # Lấy API key từ form hoặc sử dụng API key mặc định
        api_key = request.form.get('api_key', app.config['GEMINI_API_KEY'])
        
        # Cấu hình API key mới nếu có
        if api_key != app.config['GEMINI_API_KEY']:
            genai.Client(api_key=api_key)
            
        # Tạo thư mục tạm thời
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
            
        # Xóa các file cũ trong thư mục tạm
        for file in os.listdir(temp_dir):
            os.remove(os.path.join(temp_dir, file))
            
        # Lưu các file upload
        for file in files:
            if file and allowed_file(file.filename):
                file.save(os.path.join(temp_dir, file.filename))
                
        # Xử lý ảnh
        processor = SpeechBubbleProcessor(api_key)
        output_dir = processor.process_folder(temp_dir)
        
        # Tạo file zip
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], 'processed_images.zip')
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(output_dir):
                for file in files:
                    zipf.write(os.path.join(root, file), file)
                    
        return jsonify({
            'success': True,
            'message': 'Xử lý ảnh thành công',
            'download_url': url_for('download_file', filename='processed_images.zip', _external=True)
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/changeImage')
def change_image():
    return render_template('functions/changeimage.html')

@app.route('/cutmergeimage')
def cut_merge_image():
    return render_template('functions/cutmergeimage.html')

@app.route('/dowloaf')
def download_image():
    return render_template('functions/dowloaf.html')

@app.route('/mergeword')
def merge_word():
    return render_template('functions/mergeword.html')

@app.route('/renameimage')
def rename_image():
    return render_template('functions/renameimage.html')

@app.route('/ocr')
def ocr():
    return render_template('functions/ocr.html')

@app.route('/translate')
def translate():
    return render_template('functions/translate.html')

@app.route('/remove_speech_bubbles')
def remove_speech_bubbles():
    return render_template('functions/speechbubble.html')

@app.route('/test_api_key', methods=['POST'])
def test_api_key():
    try:
        data = request.json
        api_key = data.get('api_key')
        
        if not api_key:
            return jsonify({'error': 'API key không được cung cấp'})
        
        client = genai.Client(api_key=api_key)
        
        # Tạo một tin nhắn đơn giản để kiểm tra API key
        response = client.models.generate_content(
            model="gemini-2.0-flash", contents="Explain how AI works in a few words"
        )
        
        return jsonify({'success': True, 'message': 'API key hợp lệ'})
        
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/execute/download_images', methods=['POST'])
def execute_download():
    return download_selected_images()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)
